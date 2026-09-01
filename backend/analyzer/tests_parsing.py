import os
import tempfile
import docx
from django.test import TestCase
from django.conf import settings
from analyzer.services import extract_text_from_file


class ResumeParsingTests(TestCase):
    def setUp(self):
        self.temp_files = []

    def tearDown(self):
        for path in self.temp_files:
            if os.path.exists(path):
                try:
                    os.remove(path)
                except Exception:
                    pass

    def _create_temp_file(self, suffix, content, encoding=None):
        mode = 'w' if encoding else 'wb'
        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False, mode=mode, encoding=encoding) as f:
            f.write(content)
            path = f.name
        self.temp_files.append(path)
        return path

    def test_txt_parsing_utf8(self):
        content = "Jane Doe\nSkills: Python, Django, Rest Framework\n"
        path = self._create_temp_file(".txt", content.encode("utf-8"))
        text = extract_text_from_file(path, "resume.txt")
        self.assertIn("Jane Doe", text)
        self.assertIn("Python, Django", text)

    def test_txt_parsing_utf16(self):
        content = "Bob Smith\nSkills: JavaScript, React, HTML5\n"
        path = self._create_temp_file(".txt", content.encode("utf-16"))
        text = extract_text_from_file(path, "resume.txt")
        self.assertIn("Bob Smith", text)
        self.assertIn("JavaScript, React", text)

    def test_txt_parsing_latin1(self):
        content = "Alice Springs\nSkills: C++, Java, SQL\n"
        path = self._create_temp_file(".txt", content.encode("latin-1"))
        text = extract_text_from_file(path, "resume.txt")
        self.assertIn("Alice Springs", text)
        self.assertIn("C++, Java", text)

    def test_docx_parsing_paragraphs(self):
        doc = docx.Document()
        doc.add_paragraph("Charlie Brown")
        doc.add_paragraph("Skills: Ruby, Rails, Postgres")

        # Save to temp file
        path = self._create_temp_file(".docx", b"")
        doc.save(path)

        text = extract_text_from_file(path, "resume.docx")
        self.assertIn("Charlie Brown", text)
        self.assertIn("Ruby, Rails", text)

    def test_docx_parsing_tables(self):
        doc = docx.Document()
        doc.add_paragraph("David Miller")

        # Create a table for tabular resume experience/skills
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Programming Languages"
        table.cell(0, 1).text = "Go, Rust, Python"
        table.cell(1, 0).text = "Tools"
        table.cell(1, 1).text = "Kubernetes, Docker, Git"

        path = self._create_temp_file(".docx", b"")
        doc.save(path)

        text = extract_text_from_file(path, "resume.docx")
        self.assertIn("David Miller", text)
        self.assertIn("Programming Languages", text)
        self.assertIn("Go, Rust", text)
        self.assertIn("Kubernetes, Docker", text)

    def test_pdf_parsing_sample1(self):
        base_dir = getattr(settings, "BASE_DIR", "")
        possible_paths = [
            os.path.join(base_dir, "../frontend/public/sample-resume.pdf"),
            os.path.join(base_dir, "frontend/public/sample-resume.pdf"),
            "../frontend/public/sample-resume.pdf",
            "frontend/public/sample-resume.pdf",
        ]

        pdf_file_path = None
        for p in possible_paths:
            if os.path.exists(p):
                pdf_file_path = p
                break

        self.assertIsNotNone(
            pdf_file_path, "Could not find frontend/public/sample-resume.pdf in workspace")

        text = extract_text_from_file(pdf_file_path, "sample-resume.pdf")
        self.assertIn("John Doe", text)
        self.assertIn("python, django", text)

    def test_pdf_parsing_sample2_generated(self):
        base_dir = getattr(settings, "BASE_DIR", "")
        possible_paths = [
            os.path.join(base_dir, "../frontend/public/sample-resume.pdf"),
            os.path.join(base_dir, "frontend/public/sample-resume.pdf"),
            "../frontend/public/sample-resume.pdf",
            "frontend/public/sample-resume.pdf",
        ]

        pdf_file_path = None
        for p in possible_paths:
            if os.path.exists(p):
                pdf_file_path = p
                break

        self.assertIsNotNone(pdf_file_path)

        with open(pdf_file_path, "rb") as f:
            pdf_bytes = f.read()

        # Replace placeholders keeping length identical (28 and 66 chars)
        pdf2_bytes = pdf_bytes.replace(
            b"John Doe - Software Engineer",
            b"Jane Doe - Hardware Engineer"
        ).replace(
            b"Skills: python, django, react, javascript, html, css, git, github",
            b"Skills: kotlin, spring, docker, kubernetes, java, bash, aws, gitlab"
        )

        path = self._create_temp_file(".pdf", pdf2_bytes)
        text = extract_text_from_file(path, "resume2.pdf")
        self.assertIn("Jane Doe", text)
        self.assertIn("kotlin, spring", text)
