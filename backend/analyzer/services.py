import os
import pdfplumber
import docx
import textstat
from django.contrib.auth import get_user_model
from .models import ResumeAnalysis
from . import role_skills
from .skill_matcher import extract_skills, match_skills_with_partial
from .scoring import compute_score_breakdown
from .timeline import analyse as analyse_timeline
from resume_analyzer.quantify_checker import flag_unquantified_bullets

User = get_user_model()

from django.core.cache import cache

EXPERIENCE_LEVEL_SKILLS = {
    "Junior": {
        "Frontend Developer": ["html", "css", "javascript", "react", "git", "github"],
        "Backend Developer": ["python", "javascript", "sql", "git", "github", "flask", "node.js"],
        "Data Analyst": ["python", "sql", "excel", "pandas", "data analysis", "jupyter"],
    },
    "Mid-Level": {
        "Frontend Developer": [
            "html", "css", "javascript", "typescript", "react",
            "next.js", "tailwind", "git", "github", "webpack",
        ],
        "Backend Developer": [
            "python", "django", "flask", "fastapi", "node.js", "express.js",
            "sql", "mysql", "postgresql", "mongodb", "docker", "git", "github",
        ],
        "Data Analyst": [
            "python", "sql", "excel", "machine learning", "deep learning",
            "data analysis", "pandas", "numpy", "matplotlib", "tensorflow",
            "scikit-learn", "jupyter",
        ],
    },
    "Senior": {
        "Frontend Developer": [
            "html", "css", "javascript", "typescript", "react", "next.js",
            "tailwind", "git", "github", "webpack", "docker",
            "system design", "leadership", "mentoring", "ci/cd", "performance optimization",
        ],
        "Backend Developer": [
            "python", "django", "fastapi", "node.js", "postgresql", "mongodb",
            "docker", "kubernetes", "git", "github", "system design", "microservices",
            "ci/cd", "leadership", "mentoring", "distributed systems", "redis",
        ],
        "Data Analyst": [
            "python", "sql", "excel", "machine learning", "deep learning",
            "data analysis", "pandas", "numpy", "matplotlib", "tensorflow",
            "scikit-learn", "jupyter", "bigquery", "data modeling", "leadership", "mentoring",
        ],
    },
}

def get_role_skills():
    role_skills = cache.get("role_skills_dict")
    if role_skills is None:
        from .models import Role
        role_skills = {}
        for role in Role.objects.prefetch_related("skills").all():
            # Sorted because the m2m read has no ordering of its own, so two
            # calls could return the same skills in different orders. The list
            # is shown to the user and each entry is worth 1/len of the score,
            # so an unstable order makes two runs of one resume read
            # differently for no reason. Sorted in Python rather than with
            # order_by(), which would defeat the prefetch above.
            role_skills[role.name] = sorted(skill.name for skill in role.skills.all())
        cache.set("role_skills_dict", role_skills, timeout=60*60*24)
    return role_skills

def resolve_role_skills(target_role, experience_level="Mid-Level"):
    """Resolve a role's required skills, database first.

    Returns a :class:`~analyzer.role_skills.RoleSkillSet` — the list plus which
    store answered and whether the requested level was understood.

    This used to check ``EXPERIENCE_LEVEL_SKILLS`` first and only fall through
    to the database for roles the dictionary did not have. Since the dictionary
    covers every role the product ships, the database branch never ran, and
    editing a ``Role``'s skills had no effect on anything. See #708 and
    ``analyzer.role_skills`` for why the precedence is the other way round.
    """
    return role_skills.resolve(
        role=target_role,
        level=experience_level,
        database_roles=get_role_skills(),
        default_roles_by_level=EXPERIENCE_LEVEL_SKILLS,
    )


def get_role_skills_for_level(target_role, experience_level="Mid-Level"):
    """Required skills for a role at a level, as a plain list.

    Kept as the narrow form because most callers only want the list. Use
    :func:`resolve_role_skills` where the provenance matters.
    """
    return resolve_role_skills(target_role, experience_level).skills


def get_known_roles():
    """Every role either store knows about.

    ``analyze_resume`` used to build its track comparison from the database
    alone, so an unseeded ``Role`` table produced an empty comparison while the
    requested role still scored fine from the packaged defaults.
    """
    return role_skills.known_roles(get_role_skills(), EXPERIENCE_LEVEL_SKILLS)


def generate_level_tailored_suggestions(missing_skills, experience_level="Mid-Level", target_role=""):
    norm_level, _ = role_skills.normalise_level(experience_level)

    suggestions = []
    if norm_level == "Senior":
        for skill in missing_skills:
            if skill in ["leadership", "mentoring", "management"]:
                suggestions.append("Highlight engineering mentorship, cross-team alignment, and technical leadership initiatives.")
            elif skill in ["system design", "microservices", "distributed systems", "ci/cd"]:
                suggestions.append(f"Demonstrate scalable system design and architecture governance with {skill.title()}.")
            else:
                suggestions.append(f"Demonstrate senior-level production mastery, architectural decisions, and performance scaling with {skill.title()}.")
        if not any("leadership" in s.lower() or "mentorship" in s.lower() or "architectural" in s.lower() for s in suggestions):
            suggestions.append("Demonstrate senior technical leadership, architectural decision records (ADRs), and developer mentorship.")
    elif norm_level == "Junior":
        for skill in missing_skills:
            suggestions.append(f"Build foundational portfolio projects and highlight coursework or hands-on practice with {skill.title()}.")
    else:
        for skill in missing_skills:
            suggestions.append(f"Add projects or experience with {skill.title()}")

    return suggestions

PIPELINE_STAGES = [
    {"stage": "extracting", "label": "Extracting text from document", "percent": 25},
    {"stage": "matching", "label": "Detecting & matching skills", "percent": 60},
    {"stage": "scoring", "label": "Generating ATS score & recommendations", "percent": 90},
    {"stage": "done", "label": "Analysis complete", "percent": 100},
]

def calculate_readability(text):
    score = textstat.flesch_reading_ease(text)
    if score >= 60:
        label = "easy"
    elif score >= 30:
        label = "moderate"
    else:
        label = "dense"
    return round(score , 1), label

def extract_text_from_file(file_path, file_name):
    text = ""
    if file_name.lower().endswith('.docx'):
        doc = docx.Document(file_path)
        # Extract paragraph text
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        # Extract table text to capture nested tables/tabular resumes
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
    elif file_name.lower().endswith('.txt'):
        # Robust multi-encoding fallback parsing
        for encoding in ('utf-8-sig', 'utf-16', 'latin-1'):
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text = f.read()
                break
            except (UnicodeDecodeError, LookupError):
                continue
        else:
            # Final fallback in case none of the above succeeded
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()
            except Exception:
                pass
    else:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted
    return text


def analyze_cover_letter(text, target_role="", job_description=""):
    word_count = len(text.split())
    
    # 1. Length Feedback
    if word_count < 150:
        length_status = "Too short"
        length_feedback = f"Your cover letter is a bit short ({word_count} words). Consider expanding it to between 200 and 400 words to better detail your experience."
    elif word_count > 500:
        length_status = "Too long"
        length_feedback = f"Your cover letter is quite long ({word_count} words). Try to keep it concise and under 400 words so that recruiters can scan it quickly."
    else:
        length_status = "Good"
        length_feedback = f"Excellent length ({word_count} words). Your cover letter is concise and well-proportioned."
        
    # 2. Tone Feedback
    action_verbs = [
        "designed", "led", "managed", "implemented", "solved", "created", "analyzed", 
        "spearheaded", "collaborated", "executed", "developed", "built", "engineered",
        "optimized", "improved", "delivered", "facilitated", "increased", "reduced"
    ]
    weak_words = ["just", "hope", "try", "think", "believe", "maybe", "sort of", "hoping", "might"]
    
    words_lower = text.lower()
    found_action = [v for v in action_verbs if v in words_lower]
    found_weak = [w for w in weak_words if w in words_lower]
    
    tone_suggestions = []
    tone_categories = []
    
    if any(word in words_lower for word in ["excited", "thrilled", "passionate", "eager", "enthusiastic"]):
        tone_categories.append("Enthusiastic")
    if any(word in words_lower for word in ["expert", "specialist", "analytical", "background", "results", "performance"]):
        tone_categories.append("Analytical")
    
    if len(found_action) >= 4:
        tone_categories.append("Confident")
    elif len(found_action) < 2:
        tone_suggestions.append("Incorporate more active action verbs (e.g. 'spearheaded', 'implemented', 'optimized') to make your achievements sound more impactful.")
        
    if len(found_weak) > 2:
        tone_suggestions.append("Limit filler/weak words (like 'hope', 'just', 'believe') to sound more authoritative and confident in your capabilities.")
        
    tone_label = " & ".join(tone_categories) if tone_categories else "Professional"
    
    tone_feedback = f"Your cover letter has a {tone_label.lower()} tone. "
    if found_action:
        tone_feedback += f"It effectively uses active language (found action verbs: {', '.join(found_action[:3])})."
    else:
        tone_feedback += "Consider using stronger action verbs to describe your achievements."
        
    # 3. Relevance Feedback
    relevance_suggestions = []
    references_role = False
    references_company = False
    
    if target_role:
        role_words = target_role.lower().split()
        if target_role.lower() in words_lower or all(w in words_lower for w in role_words):
            references_role = True
        else:
            relevance_suggestions.append(f"Explicitly mention your target role '{target_role}' early in the letter to establish clear intent.")
            
    company_keywords = ["company", "firm", "organization", "team", "your team", "your organization", "enterprise", "agency", "institution"]
    if any(ck in words_lower for ck in company_keywords) or ("dear hiring" in words_lower or "dear team" in words_lower or "recruiting team" in words_lower):
        references_company = True
    else:
        relevance_suggestions.append("Mention the target company name or refer to their team to show that the letter is personalized.")
        
    if job_description:
        req_skills = get_role_skills().get(target_role, [])
        cover_letter_skills = [s for s in req_skills if s in words_lower]
        
        if cover_letter_skills:
            relevance_feedback = f"Great alignment! Your cover letter references several key skills required for the role, including: {', '.join([s.title() for s in cover_letter_skills[:4]])}."
        else:
            relevance_feedback = "Your cover letter does not reference the key technical skills from the target career track."
            if req_skills:
                relevance_suggestions.append(f"Weave in mentions of core role competencies like {', '.join([s.title() for s in req_skills[:3]])} to demonstrate your alignment.")
    else:
        relevance_feedback = "Please provide a job description alongside your cover letter to get specific relevance and alignment feedback."
        
    return {
        "word_count": word_count,
        "length": {
            "status": length_status,
            "feedback": length_feedback
        },
        "tone": {
            "label": tone_label,
            "feedback": tone_feedback,
            "suggestions": tone_suggestions
        },
        "relevance": {
            "references_role": references_role,
            "references_company": references_company,
            "feedback": relevance_feedback,
        }
    }
SKILL_QUESTIONS = {
    "html": [
        "What is semantic HTML, and how does it improve SEO and accessibility for assistive technologies?",
        "Explain the purpose of HTML5 data-* attributes and how they are accessed in JavaScript and CSS."
    ],
    "css": [
        "Explain the CSS Box Model, the difference between content-box and border-box, and how box-sizing affects layout sizing.",
        "Compare CSS Flexbox and Grid. In what scenarios is one layout model more appropriate than the other?"
    ],
    "javascript": [
        "Explain event delegation in JavaScript and how event bubbling enables it.",
        "Discuss JavaScript closures. Can you describe a scenario where you used closures to maintain private variables or state?"
    ],
    "typescript": [
        "How do TypeScript interfaces differ from type aliases, and when would you choose one over the other?",
        "Explain the concept of generics in TypeScript and how they help write reusable, type-safe code."
    ],
    "react": [
        "Explain React's virtual DOM reconciliation process and why key props are necessary when rendering dynamic lists.",
        "How do you manage complex state in a growing React application, and how do you prevent unnecessary child component re-renders?"
    ],
    "next.js": [
        "What are the main differences between Server-Side Rendering (SSR), Static Site Generation (SSG), and Client-Side Rendering in Next.js?",
        "Explain how Next.js App Router layout systems and nested routing improve page shell caching and state preservation."
    ],
    "tailwind": [
        "How does Tailwind's utility-first approach compare to traditional component-based CSS? Discuss performance benefits (e.g. PurgeCSS)."
    ],
    "git": [
        "Explain the difference between git merge and git rebase. Under what workflows would you choose one over the other?"
    ],
    "github": [
        "Describe your process for resolving merge conflicts during a pull request review on GitHub."
    ],
    "python": [
        "Explain the difference between Python's list comprehensions, generator expressions, and traditional loops. When is a generator preferred?",
        "How does Python's Global Interpreter Lock (GIL) affect multi-threaded CPU-bound programs, and how do you bypass it?"
    ],
    "django": [
        "How does Django's ORM handle database query optimization? How do select_related and prefetch_related prevent N+1 query patterns?",
        "Explain Django's middleware architecture. How would you design a custom middleware to log incoming request execution times?"
    ],
    "flask": [
        "Explain Flask's application context and request context. Why are they separated, and how do context locals work?"
    ],
    "fastapi": [
        "What makes FastAPI highly performant compared to standard frameworks? Discuss its dependency injection and async/await support."
    ],
    "node.js": [
        "How does Node.js handle non-blocking asynchronous I/O operations despite being single-threaded? Explain the event loop."
    ],
    "sql": [
        "How do database indexes (like B-Tree indexes) speed up query execution? Discuss the trade-offs on insert/update performance.",
        "Explain the differences between INNER JOIN, LEFT JOIN, RIGHT JOIN, and FULL OUTER JOIN with hypothetical query examples."
    ],
    "postgresql": [
        "Discuss transaction isolation levels in PostgreSQL. How does Multi-Version Concurrency Control (MVCC) prevent dirty reads?"
    ],
    "mongodb": [
        "How does schema design in a document database like MongoDB differ from a traditional relational database? When is denormalization preferred?"
    ],
    "docker": [
        "Explain the difference between Docker images and containers. How do multi-stage Docker builds optimize image size and security?"
    ],
    "excel": [
        "Describe how you handle large datasets in Excel. What formulas or features (e.g. XLOOKUP, Pivot Tables) do you use to aggregate data?"
    ],
    "machine learning": [
        "Explain the bias-variance trade-off in machine learning model training. How do regularization methods (like L1/L2) mitigate overfitting?",
        "What metrics (e.g. precision, recall, F1-score, ROC-AUC) would you prioritize when evaluating a highly imbalanced classification dataset?"
    ],
    "deep learning": [
        "Explain the vanishing gradient problem in deep neural networks and what techniques (e.g. ReLU, batch normalization, residual connections) help solve it."
    ],
    "data analysis": [
        "Describe your process for clean-up and preprocessing of a noisy, incomplete dataset before performing exploratory data analysis."
    ],
    "pandas": [
        "How do you optimize operations on large DataFrames in Pandas? Contrast the performance of apply() vs vectorized operations."
    ],
    "numpy": [
        "Explain NumPy's broadcasting rules and why vectorized calculations on NumPy arrays are faster than traditional Python loops."
    ],
}

ROLE_QUESTIONS = {
    "Frontend Developer": [
        "How do you approach web accessibility (a11y) to ensure compliance with WCAG standards in modern Single Page Applications (SPAs)?",
        "Describe how you optimize frontend performance. What metrics (e.g. LCP, FID, CLS) do you track, and what strategies do you use to improve them?",
        "How do you handle API state caching and synchronization (e.g. using React Query, RTK Query, or custom caching layers) on the frontend?"
    ],
    "Backend Developer": [
        "How do you design a secure, idempotent RESTful API endpoint? Discuss error handling schemas, validation protocols, and JWT token authentication.",
        "Describe how you would design a system architecture that handles asynchronous background task processing (e.g., using message queues like Celery, Redis, or RabbitMQ).",
        "How do you secure your backend APIs against common vulnerabilities like SQL injection, cross-site scripting (XSS), and denial-of-service (DoS) attacks?"
    ],
    "Data Analyst": [
        "Describe a time when you discovered an unexpected anomaly or insight in a dataset. How did you communicate your findings to non-technical stakeholders?",
        "How do you design dashboard KPIs that are actionable rather than just vanity metrics for executive business teams?",
        "Explain how you choose the appropriate data visualization (e.g. bar chart, scatter plot, heat map) to best tell a specific data story."
    ]
}

BEHAVIORAL_QUESTIONS = [
    "Tell me about a challenging technical bug you encountered in a recent project. How did you diagnose and resolve it?",
    "Describe a situation where you had to work with a teammate who had a different technical opinion. How did you reach a consensus?",
    "How do you keep up with new technology changes, libraries, and framework updates in your engineering track?",
    "Describe a project where you had a tight deadline and limited requirements. How did you prioritize tasks to deliver value on time?"
]

def generate_interview_questions(skills, target_role):
    import random
    questions = []
    
    skills_shuffled = list(skills)
    random.shuffle(skills_shuffled)
    
    for skill in skills_shuffled:
        skill_lower = skill.lower()
        if skill_lower in SKILL_QUESTIONS:
            q = random.choice(SKILL_QUESTIONS[skill_lower])
            if q not in questions:
                questions.append(q)
            if len(questions) >= 4:
                break
                
    role_pool = ROLE_QUESTIONS.get(target_role, [])
    if role_pool:
        role_shuffled = list(role_pool)
        random.shuffle(role_shuffled)
        for q in role_shuffled:
            if q not in questions:
                questions.append(q)
            if len(questions) >= 6:
                break
                
    behavioral_shuffled = list(BEHAVIORAL_QUESTIONS)
    random.shuffle(behavioral_shuffled)
    for q in behavioral_shuffled:
        if len(questions) >= 8:
            break
        if q not in questions:
            questions.append(q)
            
    if len(questions) < 5:
        for q in behavioral_shuffled:
            if q not in questions:
                questions.append(q)
            if len(questions) >= 5:
                break
                
    return questions[:8]


def analyze_resume(file_path, target_role, file_name="resume.pdf", user_id=None, job_description=None, cover_letter_path=None, cover_letter_name=None, experience_level="Mid-Level"):
    text = ""
    try:
        text = extract_text_from_file(file_path, file_name)
    finally:
        if os.path.exists(file_path):
            os.remove(file_path)

    raw_text = text
    readability_score, readability_label = calculate_readability(raw_text)
    detected = extract_skills(text)

    if job_description and job_description.strip():
        required = extract_skills(job_description)
        role_skill_set = None
    else:
        role_skill_set = resolve_role_skills(target_role, experience_level)
        required = role_skill_set.skills

    matched, partial, missing = match_skills_with_partial(required, raw_text, detected)

    score_credit = len(matched) + (0.5 * len(partial))
    score = (
        int(score_credit / len(required) * 100)
        if required
        else min(len(detected) * 10, 100)
    )

    suggestions = []
    for item in partial:
        suggestions.append(f"Near match: Your resume mentions '{item['matched_variant']}' which is a partial match for target skill '{item['skill']}'. Clarify or explicitly list '{item['skill']}' for full credit.")

    suggestions.extend(generate_level_tailored_suggestions(missing, experience_level, target_role))

    # Process optional cover letter if provided
    cover_letter_text = ""
    cover_letter_feedback = None
    if cover_letter_path:
        try:
            cover_letter_text = extract_text_from_file(cover_letter_path, cover_letter_name)
            cover_letter_feedback = analyze_cover_letter(cover_letter_text, target_role, job_description)
        finally:
            if os.path.exists(cover_letter_path):
                os.remove(cover_letter_path)

    # Generate interview questions based on target track and detected skills
    interview_questions = generate_interview_questions(detected, target_role)

    analysis_id = None

    if user_id:
        try:
            user = User.objects.get(id=user_id)
            analysis_record = ResumeAnalysis.objects.create(
                user=user,
                file_name=file_name,
                target_role=target_role,
                experience_level=experience_level or "Mid-Level",
                job_description=job_description,
                score=score,
                skills_found=detected,
                suggestions=suggestions,
                matched_skills=matched,
                partial_skills=partial,
                missing_skills=missing,
                resume_text=raw_text,
                cover_letter_text=cover_letter_text if cover_letter_text else None,
                cover_letter_feedback=cover_letter_feedback,
                interview_questions=interview_questions,
            )
            analysis_id = analysis_record.id
        except User.DoesNotExist:
            pass

    progress_info = {
        "current_stage": "done",
        "percent": 100,
        "stages": PIPELINE_STAGES,
    }

    # Every role either store knows about, not just the database's. An unseeded
    # `Role` table used to produce an empty comparison table while the requested
    # role scored perfectly well from the packaged defaults -- two different
    # answers to "which roles exist" in one response. See #708.
    track_comparisons = {}
    for role in get_known_roles():
        role_set = resolve_role_skills(role, experience_level)
        role_req_skills = role_set.skills
        role_matched, role_partial, role_missing = match_skills_with_partial(role_req_skills, raw_text, detected)
        role_credit = len(role_matched) + (0.5 * len(role_partial))
        role_score = (
            int(role_credit / len(role_req_skills) * 100)
            if role_req_skills
            else min(len(detected) * 10, 100)
        )
        role_suggestions = generate_level_tailored_suggestions(role_missing, experience_level, role)

        track_comparisons[role] = {
            "score": role_score,
            "matched_skills": role_matched,
            "partial_skills": role_partial,
            "missing_skills": role_missing,
            "suggestions": role_suggestions,
            # Which store answered for this row. Two rows in one table can come
            # from different sources, and a score whose provenance is invisible
            # is one nobody can debug.
            "skills_source": role_set.source,
        }

    quantify_nudges = flag_unquantified_bullets(raw_text.split('\n'))

    # Employment timeline. Deliberately kept out of `score` for the same reason
    # `score_breakdown` is: that number is persisted on ResumeAnalysis and read
    # by the leaderboard, version comparison and the digest, so changing what it
    # means would change every historical row. See #709.
    resume_timeline = analyse_timeline(raw_text, experience_level=experience_level)

    # Multi-factor view of the same resume. `score` above stays the keyword
    # ratio — it is persisted on ResumeAnalysis and read by the leaderboard,
    # version comparison and digest, so redefining it would change the meaning
    # of every historical row.
    score_breakdown = compute_score_breakdown(
        text=raw_text,
        matched_skills=matched,
        required_skills=required,
        detected_skills=detected,
        readability_score=readability_score,
        readability_label=readability_label,
        quantify_nudges=quantify_nudges,
        partial_skills=partial,
    )

    return {
        "id": analysis_id,
        "score": score,
        "score_breakdown": score_breakdown.as_dict(),
        "readability_score": readability_score,
        "readability_label": readability_label,
        "skills_found": detected,
        "suggestions": suggestions,
        "quantify_nudges": quantify_nudges,
        "timeline": resume_timeline.as_dict(),
        "matched_skills": matched,
        "partial_skills": partial,
        "missing_skills": missing,
        "target_role": target_role,
        "experience_level": experience_level or "Mid-Level",
        # What the scoring actually used, and where the requirements came from.
        # The response used to echo back whatever level the caller sent while
        # scoring against a different one -- "Principal" was silently read as
        # Mid-Level and nothing said so. See #708.
        "role_skills": role_skill_set.as_dict() if role_skill_set else {
            "role": target_role,
            "level": role_skills.normalise_level(experience_level)[0],
            "level_as_requested": experience_level if isinstance(experience_level, str) else "",
            "level_recognised": role_skills.normalise_level(experience_level)[1],
            "skills": list(required),
            # A pasted job description overrides both stores, which is the
            # documented behaviour and worth naming rather than leaving the
            # caller to infer it from an absent field.
            "source": "job-description",
        },
        "resume_text": raw_text,
        "cover_letter_text": cover_letter_text if cover_letter_text else None,
        "cover_letter_feedback": cover_letter_feedback,
        "interview_questions": interview_questions,
        "progress": progress_info,
        "pipeline_stages": PIPELINE_STAGES,
        "track_comparisons": track_comparisons,
    }