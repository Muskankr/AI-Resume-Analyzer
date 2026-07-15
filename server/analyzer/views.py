from rest_framework.decorators import api_view, parser_classes, permission_classes, throttle_classes
from rest_framework.parsers import MultiPartParser, FormParser
from rest_framework.permissions import IsAuthenticated, AllowAny
from rest_framework.response import Response
from rest_framework.throttling import SimpleRateThrottle
from rest_framework import status
import os

import pdfplumber
from docx import Document

from django.conf import settings

from .models import ResumeAnalysis
from .serializers import SignupSerializer, ResumeAnalysisSerializer


# Supported upload formats — kept in one place so frontend/backend stay in sync.
SUPPORTED_RESUME_EXTENSIONS = (".pdf", ".docx")


def extract_text_from_file(file_obj):
    """
    Extract plain text from an uploaded resume file.

    Dispatches to the right parser based on the file extension so that the
    rest of the analysis pipeline (skill detection, scoring, suggestions)
    can stay 100% format-agnostic and produce the same output shape
    regardless of whether the source was a PDF or a Word document.

    Args:
        file_obj: A Django UploadedFile (has .name and a file-like .file).

    Returns:
        str: The lowercased text content of the resume.

    Raises:
        ValueError: If the file extension is not supported or parsing fails.
    """
    original_name = getattr(file_obj, "name", "") or ""
    ext = os.path.splitext(original_name)[1].lower()

    if ext == ".pdf":
        text = ""
        # pdfplumber needs to read from the underlying file pointer.
        # Reset to start in case anything has touched it already.
        if hasattr(file_obj, "seek"):
            file_obj.seek(0)
        with pdfplumber.open(file_obj) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text

    if ext == ".docx":
        # python-docx reads from the .docx zip archive directly.
        if hasattr(file_obj, "seek"):
            file_obj.seek(0)
        document = Document(file_obj)
        # Paragraphs cover most resume content; tables are also common in
        # Word resumes (e.g. contact info, experience grids), so include them.
        parts = [p.text for p in document.paragraphs if p.text]
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        parts.append(cell.text)
        return "\n".join(parts)

    raise ValueError(
        f"Unsupported file type '{ext or 'unknown'}'. "
        f"Supported types: {', '.join(SUPPORTED_RESUME_EXTENSIONS)}"
    )


class UploadRateThrottle(SimpleRateThrottle):
    scope = "upload"

    def get_rate(self):
        return getattr(settings, "RESUME_UPLOAD_RATE", "10/hour")

    def get_cache_key(self, request, view):
        ident = self.get_ident(request)  # client IP
        return self.cache_format % {"scope": self.scope, "ident": ident}

skills_list = [
    "python", "django", "react", "javascript", "sql",
    "html", "css", "git", "github", "flask",
    "machine learning", "data analysis",
    "excel", "microsoft office", "ms office",
    "c", "c++", "java"
]

ROLE_SKILL_MATRICES = {
    "Frontend Developer": ["html", "css", "javascript", "react", "git", "github"],
    "Backend Developer": ["python", "django", "flask", "sql", "git", "github"],
    "Data Analyst": ["python", "excel", "sql", "data analysis", "machine learning"]
}


@api_view(["POST"])
@permission_classes([AllowAny])
def signup(request):
    serializer = SignupSerializer(data=request.data)
    if serializer.is_valid():
        serializer.save()
        return Response({"detail": "Account created."}, status=status.HTTP_201_CREATED)
    return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


@api_view(["POST"])
@parser_classes([MultiPartParser, FormParser])
@permission_classes([AllowAny])
@throttle_classes([UploadRateThrottle])
def upload_resume(request):
    file = request.FILES.get("file")
    target_role = request.data.get("role", None)

    # --- Input validation -------------------------------------------------
    if not file:
        return Response(
            {"detail": "No file uploaded. Please attach a .pdf or .docx resume."},
            status=status.HTTP_400_BAD_REQUEST,
        )

    file_name = file.name
    ext = os.path.splitext(file_name)[1].lower()
    if ext not in SUPPORTED_RESUME_EXTENSIONS:
        return Response(
            {
                "detail": (
                    f"Unsupported file type '{ext or 'unknown'}'. "
                    f"Please upload one of: {', '.join(SUPPORTED_RESUME_EXTENSIONS)}."
                )
            },
            status=status.HTTP_400_BAD_REQUEST,
        )

    # --- Text extraction (format-agnostic from here on) -------------------
    try:
        text = extract_text_from_file(file)
    except Exception as exc:
        # Unsupported extension, corrupted file, encrypted PDF, malformed
        # .docx, etc. — all surface to the client as a 400 with a helpful
        # message so the rest of the API contract stays stable.
        return Response(
            {"detail": str(exc) or "Could not read the uploaded file."},
            status=status.HTTP_400_BAD_REQUEST,
        )

    text = text.lower()

    detected_skills = [s for s in skills_list if s.lower() in text]

    matched_skills = []
    missing_skills = []

    if target_role in ROLE_SKILL_MATRICES:
        required_skills = ROLE_SKILL_MATRICES[target_role]

        for skill in required_skills:
            if skill in detected_skills:
                matched_skills.append(skill)
            else:
                missing_skills.append(skill)

        # Dynamic role-based score
        score = int((len(matched_skills) / len(required_skills)) * 100) if required_skills else 100

        # Dynamic suggestions
        suggestions = [
            f"Add experience or projects with {skill.upper() if skill in ['html', 'css', 'sql', 'git'] else skill.capitalize()}"
            for skill in missing_skills
        ]
    else:
        score = min(len(detected_skills) * 10, 100)

        suggestions = []
        if "python" not in detected_skills:
            suggestions.append("Add Python projects")
        if "django" not in detected_skills:
            suggestions.append("Mention Django experience")
        if "react" not in detected_skills:
            suggestions.append("Add frontend skills like React")

    # Save to DB only for authenticated users
    if request.user and request.user.is_authenticated:
        ResumeAnalysis.objects.create(
            user=request.user,
            file_name=file_name,
            score=score,
            skills_found=detected_skills,
            suggestions=suggestions,
            matched_skills=matched_skills,
            missing_skills=missing_skills,
            target_role=target_role or "",
        )

    return Response({
        "score": score,
        "skills_found": detected_skills,
        "suggestions": suggestions,
        "target_role": target_role,
        "matched_skills": matched_skills,
        "missing_skills": missing_skills,
    })


@api_view(["GET"])
@permission_classes([IsAuthenticated])
def analysis_history(request):
    analyses = ResumeAnalysis.objects.filter(user=request.user)
    serializer = ResumeAnalysisSerializer(analyses, many=True)
    return Response(serializer.data)
